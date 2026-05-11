from docx import Document
from PIL import Image
import pytesseract
from transformers import pipeline

contract_qa = None

questions_per_field = {
    "agreement_value":      "What is the monthly rent amount?",
    "agreement_start_date": "When does the rental agreement start?",
    "agreement_end_date":   "When does the rental agreement end or expire?",
    "renewal_notice_days":  "How many days of notice are required to terminate or renew the agreement?",
    "party_one":            "Who is the owner or lessor of the property?",
    "party_two":            "Who is the tenant or lessee?",
}


def load_model(model_path="fine_tuned_contract_qa"):
    global contract_qa
    contract_qa = pipeline(
        "question-answering",
        model=model_path,
        tokenizer="Rakib/roberta-base-on-cuad",
    )
    print(f"Model loaded from {model_path}")


def extract_text_from_docx(docx_path):
    document = Document(docx_path)
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for table_row in table.rows:
            for table_cell in table_row.cells:
                if table_cell.text.strip():
                    parts.append(table_cell.text)
    return "\n".join(parts)


def extract_text_from_png(image_path):
    image = Image.open(image_path).convert("RGB")
    return pytesseract.image_to_string(image)


def extract_text(file_path):
    if file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    if file_path.endswith(".png"):
        return extract_text_from_png(file_path)
    raise ValueError(f"Unsupported file type: {file_path}")


def get_best_answer_for_question(full_text, question):
    """
    Run the QA pipeline once on the entire context. The pipeline natively
    handles long inputs by sliding a window of `max_seq_len` tokens with
    `doc_stride` overlap, then returns the highest-scoring span across all
    windows. This replaces the previous word-based chunking which split the
    text in arbitrary places and could miss answers that crossed a chunk
    boundary.
    """
    if not full_text or not full_text.strip():
        return ""
    try:
        result = contract_qa(
            question=question,
            context=full_text,
            max_seq_len=384,
            doc_stride=128,
            handle_impossible_answer=False,
            top_k=1,
        )
        return result["answer"].strip()
    except Exception as error:
        print(f"  QA failed: {error}")
        return ""


def extract_fields_from_text(full_text):
    extracted_fields = {}
    for field_name, question in questions_per_field.items():
        extracted_fields[field_name] = get_best_answer_for_question(
            full_text, question)
    return extracted_fields
